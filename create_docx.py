"""
Script to create Business Logic Guide DOCX document in Russian
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_document():
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Title style
    title_style = styles['Title']
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 51, 102)
    
    # Heading 1
    h1_style = styles['Heading 1']
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 102, 153)
    
    # Heading 2
    h2_style = styles['Heading 2']
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(51, 102, 153)
    
    # ==================== TITLE PAGE ====================
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title_run = title.add_run("📘 Sellary POS")
    title_run.bold = True
    title_run.font.size = Pt(36)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Руководство по бизнес-логике")
    subtitle_run.font.size = Pt(24)
    subtitle_run.font.color.rgb = RGBColor(51, 102, 153)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.add_run("Для продукт-менеджеров и клиентов").font.size = Pt(14)
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    meta_info = doc.add_paragraph()
    meta_info.add_run("Версия: ").bold = True
    meta_info.add_run("1.0\n")
    meta_info.add_run("Дата обновления: ").bold = True
    meta_info.add_run("1 февраля 2026\n")
    meta_info.add_run("Целевая аудитория: ").bold = True
    meta_info.add_run("Продукт-менеджеры, бизнес-аналитики, клиенты")
    meta_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ==================== TABLE OF CONTENTS ====================
    doc.add_heading("📋 Оглавление", level=1)
    
    toc_items = [
        "1. Краткий обзор системы",
        "2. Основные бизнес-сущности",
        "3. Продажи и касса",
        "4. Управление складом",
        "5. Поставщики и закупки",
        "6. Роли пользователей и доступ",
        "7. Финансовые расчёты",
        "8. Отчёты и аналитика",
        "9. Часто задаваемые вопросы",
        "10. Глоссарий"
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # ==================== SECTION 1: OVERVIEW ====================
    doc.add_heading("📌 1. Краткий обзор системы", level=1)
    
    doc.add_paragraph(
        "Sellary POS — современная система для точек продаж, разработанная для малого и среднего "
        "розничного бизнеса. Система оптимизирует ежедневные операции магазина."
    )
    
    doc.add_heading("Ключевые возможности", level=2)
    
    # Create capabilities table
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "Возможность"
    headers[1].text = "Ценность для бизнеса"
    
    data = [
        ("Быстрое оформление продаж", "Сокращение времени обслуживания клиентов благодаря сканированию штрих-кодов"),
        ("Учёт товаров в реальном времени", "Всегда актуальные остатки; невозможно продать больше, чем есть"),
        ("Управление поставщиками", "Отслеживание закупок и автоматизация заказов"),
        ("Информативные отчёты", "Принимайте решения на основе данных о прибыли и продажах"),
        ("Разграничение доступа", "Безопасный доступ для администраторов, менеджеров и кассиров")
    ]
    
    for i, (cap, value) in enumerate(data, 1):
        row = table.rows[i].cells
        row[0].text = cap
        row[1].text = value
    
    doc.add_paragraph()
    
    doc.add_heading("🎯 Целевые пользователи", level=2)
    users = [
        "Владельцы розничных магазинов — модернизация бизнес-процессов",
        "Менеджеры магазинов — контроль склада и продаж",
        "Кассиры — оформление ежедневных транзакций"
    ]
    for user in users:
        doc.add_paragraph(user, style='List Bullet')
    
    doc.add_page_break()
    
    # ==================== SECTION 2: ENTITIES ====================
    doc.add_heading("🏷️ 2. Основные бизнес-сущности", level=1)
    
    doc.add_heading("Описание сущностей", level=2)
    
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "Сущность"
    headers[1].text = "Описание"
    headers[2].text = "Ключевые атрибуты"
    
    entities = [
        ("Пользователь", "Сотрудник, работающий с системой", "Логин, Роль, Статус"),
        ("Товар", "Единица продажи", "Штрих-код, Название, Себестоимость, Цена, Остаток"),
        ("Категория", "Группировка товаров", "Название, Описание"),
        ("Поставщик", "Компания-поставщик товаров", "Название, Контакты, Условия оплаты"),
        ("Продажа", "Завершённая транзакция (чек)", "Сумма, Способ оплаты, Кассир, Дата"),
        ("Заказ поставщику", "Заявка на пополнение склада", "Поставщик, Позиции, Статус"),
        ("Журнал склада", "История всех движений товара", "Товар, Изменение, Причина, Пользователь")
    ]
    
    for i, (entity, desc, attrs) in enumerate(entities, 1):
        row = table.rows[i].cells
        row[0].text = entity
        row[1].text = desc
        row[2].text = attrs
    
    doc.add_page_break()
    
    # ==================== SECTION 3: SALES ====================
    doc.add_heading("💰 3. Продажи и касса", level=1)
    
    doc.add_heading("Жизненный цикл продажи", level=2)
    
    steps = [
        "1. ДОБАВИТЬ В КОРЗИНУ — Сканирование штрих-кода или поиск по имени, проверка остатка",
        "2. ПРОВЕРИТЬ КОРЗИНУ — Изменить количество, удалить товар, просмотр итогов",
        "3. ОФОРМИТЬ ОПЛАТУ — Выбрать способ оплаты, завершить продажу"
    ]
    for step in steps:
        doc.add_paragraph(step)
    
    doc.add_heading("Правила обработки продаж", level=2)
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "Правило"
    headers[1].text = "Описание"
    headers[2].text = "Влияние на бизнес"
    
    rules = [
        ("Проверка остатков", "Нельзя добавить больше, чем есть на складе", "Защита от пересортицы"),
        ("Атомарная транзакция", "Все операции выполняются вместе или не выполняются", "Целостность данных"),
        ("Автоматическое списание", "Остаток уменьшается сразу после продажи", "Точность учёта"),
        ("Аудит", "Каждая продажа создаёт записи в журнале склада", "Полная прослеживаемость")
    ]
    
    for i, (rule, desc, impact) in enumerate(rules, 1):
        row = table.rows[i].cells
        row[0].text = rule
        row[1].text = desc
        row[2].text = impact
    
    doc.add_paragraph()
    
    doc.add_heading("Поддерживаемые способы оплаты", level=2)
    
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "Способ"
    headers[1].text = "Варианты"
    headers[2].text = "Дополнительно"
    
    payments = [
        ("💵 Наличные", "Стандартная оплата", "Без дополнительных полей"),
        ("💳 Карта", "Алиф Банк, Эсхата Банк, DC", "Тип карты сохраняется для отчётов"),
        ("📱 Мобильный платёж", "Мобильные платёжные системы", "Готовность к интеграции")
    ]
    
    for i, (method, options, info) in enumerate(payments, 1):
        row = table.rows[i].cells
        row[0].text = method
        row[1].text = options
        row[2].text = info
    
    doc.add_paragraph()
    
    doc.add_heading("Статусы продажи", level=2)
    
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "Статус"
    headers[1].text = "Значение"
    headers[2].text = "Влияние на склад"
    headers[3].text = "Можно отменить?"
    
    statuses = [
        ("Завершена", "Обычная продажа", "Товар списан", "Да"),
        ("Отменена", "Продажа аннулирована", "Товар возвращён", "Нет (конечный статус)"),
        ("Частичный возврат", "Некоторые товары возвращены", "Частичный возврат", "Можно завершить"),
        ("Возврат", "Все товары возвращены", "Полный возврат", "Нет (конечный статус)")
    ]
    
    for i, (status, meaning, stock, undo) in enumerate(statuses, 1):
        row = table.rows[i].cells
        row[0].text = status
        row[1].text = meaning
        row[2].text = stock
        row[3].text = undo
    
    doc.add_paragraph()
    
    doc.add_heading("Мульти-вкладки (Параллельные сессии)", level=2)
    doc.add_paragraph(
        "Бизнес-потребность: Кассиру часто нужно приостановить обслуживание одного клиента "
        "для обслуживания другого (клиент ищет кошелёк, проверка цены и т.д.)"
    )
    
    features = [
        "Кассир может открыть несколько сессий корзины (вкладок)",
        "Переключение между вкладками без потери товаров",
        "Каждая вкладка сохраняет своё состояние корзины",
        "Продажи оформляются независимо по каждой вкладке"
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')
    
    doc.add_heading("Горячие клавиши", level=2)
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "Клавиша"
    headers[1].text = "Действие"
    
    hotkeys = [
        ("F2", "Открыть поиск товаров"),
        ("Enter", "Завершить продажу (когда открыто окно оплаты)")
    ]
    
    for i, (key, action) in enumerate(hotkeys, 1):
        row = table.rows[i].cells
        row[0].text = key
        row[1].text = action
    
    doc.add_page_break()
    
    # ==================== SECTION 4: INVENTORY ====================
    doc.add_heading("📦 4. Управление складом", level=1)
    
    doc.add_heading("Правила движения товаров", level=2)
    
    doc.add_paragraph().add_run("Приход (Увеличение остатка):").bold = True
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "Источник"
    headers[1].text = "Триггер"
    headers[2].text = "Описание"
    
    inbound = [
        ("Приёмка заказа", "Отметить товары как полученные", "Товар прибыл от поставщика"),
        ("Отмена продажи", "Отменить завершённую продажу", "Возврат товара"),
        ("Возврат от покупателя", "Оформить возврат", "Клиент вернул товар"),
        ("Ручная корректировка", "Действие администратора", "Пересчёт, найденный товар")
    ]
    
    for i, (src, trigger, desc) in enumerate(inbound, 1):
        row = table.rows[i].cells
        row[0].text = src
        row[1].text = trigger
        row[2].text = desc
    
    doc.add_paragraph()
    doc.add_paragraph().add_run("Расход (Уменьшение остатка):").bold = True
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "Источник"
    headers[1].text = "Триггер"
    headers[2].text = "Описание"
    
    outbound = [
        ("Завершение продажи", "Оформление чека", "Товары проданы клиенту"),
        ("Ручная корректировка", "Действие администратора", "Бой, недостача, образцы")
    ]
    
    for i, (src, trigger, desc) in enumerate(outbound, 1):
        row = table.rows[i].cells
        row[0].text = src
        row[1].text = trigger
        row[2].text = desc
    
    doc.add_paragraph()
    
    doc.add_heading("Оповещения о низком остатке", level=2)
    
    doc.add_paragraph("Как это работает:")
    steps = [
        "Для каждого товара задаётся Минимальный остаток (настраивается)",
        "Система постоянно сравнивает текущий остаток с минимальным",
        "Товары с остатком ниже минимума отображаются в предупреждениях на главной панели"
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Пример:\n").bold = True
    p.add_run('Товар: "Премиум Кофе 500г"\n')
    p.add_run("Текущий остаток: 3 шт.\n")
    p.add_run("Минимальный остаток: 10 шт.\n")
    p.add_run('→ ⚠️ Предупреждение: "Премиум Кофе 500г требует пополнения!"')
    
    doc.add_heading("Журнал склада (Аудит)", level=2)
    
    doc.add_paragraph("Каждое изменение остатка фиксируется навсегда:")
    
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "Поле"
    headers[1].text = "Назначение"
    
    fields = [
        ("Товар", "Какой товар был затронут"),
        ("Изменение количества", "+5, -2 и т.д."),
        ("Предыдущий остаток", "Остаток до изменения"),
        ("Новый остаток", "Остаток после изменения"),
        ("Причина", '"Продажа #1234", "Корректировка"'),
        ("Пользователь", "Кто внёс изменение"),
        ("Дата и время", "Когда это произошло")
    ]
    
    for i, (field, purpose) in enumerate(fields, 1):
        row = table.rows[i].cells
        row[0].text = field
        row[1].text = purpose
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Ценность для бизнеса: ").bold = True
    p.add_run("Полная прозрачность и ответственность при расхождениях в инвентаризации.")
    
    doc.add_page_break()
    
    # ==================== SECTION 5: PROCUREMENT ====================
    doc.add_heading("🚚 5. Поставщики и закупки", level=1)
    
    doc.add_heading("Управление поставщиками", level=2)
    
    doc.add_paragraph("Поставщики — это компании, у которых вы закупаете товары для продажи.")
    
    doc.add_paragraph().add_run("Отслеживаемая информация:").bold = True
    info = [
        "Название компании и контактное лицо",
        "Телефон, email, адрес",
        'Условия оплаты (например, "Нетто 30", "Предоплата", "При доставке")'
    ]
    for item in info:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("Жизненный цикл заказа поставщику", level=2)
    
    doc.add_paragraph("Заказы поставщику отслеживают процесс пополнения склада:")
    
    workflow = [
        "ЧЕРНОВИК — Создание/редактирование заказа",
        "ОТПРАВЛЕН — Заказ отправлен поставщику",
        "ЧАСТИЧНО ПОЛУЧЕН — Ожидаются ещё товары",
        "ПОЛУЧЕН — Все товары получены (конечный статус)",
        "ОТМЕНЁН — Заказ отменён (конечный статус)"
    ]
    for step in workflow:
        doc.add_paragraph(step)
    
    doc.add_heading("Матрица действий по статусам заказа", level=2)
    
    table = doc.add_table(rows=6, cols=6)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "Действие"
    headers[1].text = "Черновик"
    headers[2].text = "Отправлен"
    headers[3].text = "Частично получен"
    headers[4].text = "Получен"
    headers[5].text = "Отменён"
    
    actions = [
        ("Редактировать", "✅", "❌", "❌", "❌", "❌"),
        ("Отправить", "✅", "❌", "❌", "❌", "❌"),
        ("Принять товары", "❌", "✅", "✅", "❌", "❌"),
        ("Отменить", "✅", "✅", "✅", "❌", "❌"),
        ("Удалить", "✅", "❌", "❌", "❌", "❌")
    ]
    
    for i, (action, *statuses) in enumerate(actions, 1):
        row = table.rows[i].cells
        row[0].text = action
        for j, status in enumerate(statuses):
            row[j+1].text = status
    
    doc.add_paragraph()
    
    doc.add_heading("Приёмка товаров (Поддержка частичной поставки)", level=2)
    
    doc.add_paragraph().add_run("Бизнес-сценарий: ").bold = True
    doc.add_paragraph("Поставщик привёз 50 единиц, а вы заказывали 100.")
    
    doc.add_paragraph().add_run("Как это работает:").bold = True
    steps = [
        'Откройте "Принять товары" для заказа',
        "Введите фактически полученное количество по каждой позиции",
        "Система обновит: остаток товара, счётчик получено, статус заказа",
        "Создаются записи в журнале склада со ссылкой на заказ"
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")
    
    doc.add_page_break()
    
    # ==================== SECTION 6: ROLES ====================
    doc.add_heading("👥 6. Роли пользователей и доступ", level=1)
    
    doc.add_heading("Иерархия ролей", level=2)
    
    roles_desc = [
        ("АДМИНИСТРАТОР", "Полный доступ к системе, управление пользователями"),
        ("МЕНЕДЖЕР", "Отчёты, склад, контроль продаж"),
        ("КАССИР", "Работа с кассой, просмотр товаров")
    ]
    
    for role, desc in roles_desc:
        p = doc.add_paragraph()
        p.add_run(role + ": ").bold = True
        p.add_run(desc)
    
    doc.add_heading("Матрица прав доступа", level=2)
    
    table = doc.add_table(rows=13, cols=4)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "Возможность"
    headers[1].text = "Админ"
    headers[2].text = "Менеджер"
    headers[3].text = "Кассир"
    
    permissions = [
        ("Оформлять продажи (POS)", "✅", "✅", "✅"),
        ("Просматривать свои продажи", "✅", "✅", "✅"),
        ("Просматривать все продажи", "✅", "✅", "❌"),
        ("Отменять продажи", "✅", "✅", "❌"),
        ("Управлять товарами", "✅", "✅", "❌"),
        ("Просматривать склад", "✅", "✅", "🔍"),
        ("Корректировать остатки", "✅", "✅", "❌"),
        ("Управлять поставщиками", "✅", "✅", "❌"),
        ("Создавать/управлять заказами", "✅", "✅", "❌"),
        ("Просматривать отчёты", "✅", "✅", "❌"),
        ("Управлять пользователями", "✅", "❌", "❌"),
        ("Настройки системы", "✅", "❌", "❌")
    ]
    
    for i, (cap, admin, manager, cashier) in enumerate(permissions, 1):
        row = table.rows[i].cells
        row[0].text = cap
        row[1].text = admin
        row[2].text = manager
        row[3].text = cashier
    
    doc.add_page_break()
    
    # ==================== SECTION 7: CALCULATIONS ====================
    doc.add_heading("🧮 7. Финансовые расчёты", level=1)
    
    doc.add_heading("Расчёт цен и налогов", level=2)
    
    doc.add_paragraph("Система использует точные десятичные вычисления для всех финансовых операций.")
    
    doc.add_paragraph().add_run("Для каждой позиции продажи:").bold = True
    formulas = [
        "Подытог = Количество × Цена за единицу",
        "Сумма налога = Подытог × (Процент налога ÷ 100)",
        "Итого по позиции = Подытог + Сумма налога - Скидка"
    ]
    for f in formulas:
        doc.add_paragraph(f, style='List Bullet')
    
    doc.add_paragraph().add_run("Для всей продажи (чека):").bold = True
    formulas = [
        "Подытог продажи = Сумма подытогов всех позиций",
        "Налог продажи = Сумма налогов всех позиций",
        "Итого продажи = Подытог + Налог - Скидка на продажу"
    ]
    for f in formulas:
        doc.add_paragraph(f, style='List Bullet')
    
    doc.add_heading("Расчёт прибыли", level=2)
    
    doc.add_paragraph().add_run("Для каждого товара:").bold = True
    formulas = [
        "Прибыль = (Цена продажи - Себестоимость) × Проданное количество",
        "Маржа (%) = ((Цена продажи - Себестоимость) ÷ Себестоимость) × 100"
    ]
    for f in formulas:
        doc.add_paragraph(f, style='List Bullet')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Пример:\n").bold = True
    p.add_run('Товар: "Чехол для ноутбука"\n')
    p.add_run("Себестоимость: 150.00 сомони\n")
    p.add_run("Цена продажи: 250.00 сомони\n")
    p.add_run("Продано: 10 шт.\n\n")
    p.add_run("Расчёты:\n")
    p.add_run("• Прибыль с единицы: 250.00 - 150.00 = 100.00 сомони\n")
    p.add_run("• Общая прибыль: 100.00 × 10 = 1000.00 сомони\n")
    p.add_run("• Маржа: (100.00 ÷ 150.00) × 100 = 66.67%")
    
    doc.add_heading("Оценка стоимости склада", level=2)
    
    doc.add_paragraph("Общая стоимость склада рассчитывается как:")
    p = doc.add_paragraph()
    p.add_run("Стоимость склада = Σ (Остаток товара × Себестоимость товара)").italic = True
    
    doc.add_paragraph("Это представляет себестоимость вашего текущего склада.")
    
    doc.add_page_break()
    
    # ==================== SECTION 8: REPORTS ====================
    doc.add_heading("📊 8. Отчёты и аналитика", level=1)
    
    doc.add_heading("Виджеты главной панели", level=2)
    
    doc.add_paragraph("Панель управления показывает ключевые бизнес-показатели:")
    
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "Виджет"
    headers[1].text = "Отображаемые данные"
    headers[2].text = "Обновление"
    
    widgets = [
        ("Продажи сегодня", "Общая выручка за текущий день", "В реальном времени"),
        ("Прибыль сегодня", "Выручка минус себестоимость за сегодня", "В реальном времени"),
        ("Количество чеков", "Число продаж сегодня", "В реальном времени"),
        ("Низкий остаток", "Товары ниже минимального остатка", "В реальном времени"),
        ("Топ продаж", "Топ-5 по количеству за сегодня", "В реальном времени"),
        ("Последние продажи", "10 последних транзакций", "В реальном времени")
    ]
    
    for i, (widget, data, update) in enumerate(widgets, 1):
        row = table.rows[i].cells
        row[0].text = widget
        row[1].text = data
        row[2].text = update
    
    doc.add_paragraph()
    
    doc.add_heading("Доступные отчёты", level=2)
    
    reports = [
        ("📈 Отчёт по продажам", "Отслеживание трендов выручки", "Дневные продажи за выбранный период", "Линейный график дневной динамики", "7 дней, 30 дней, 90 дней, 1 год"),
        ("💵 Отчёт по прибыли", "Понимание рентабельности", "Выручка, себестоимость, прибыль, маржа", "Столбчатая диаграмма сравнения", "7 дней, 30 дней, 90 дней, 1 год"),
        ("🏆 Топ товаров", "Определение бестселлеров", "Товары по количеству продаж", "Горизонтальная столбчатая диаграмма", "7 дней, 30 дней, 90 дней, 1 год")
    ]
    
    for name, purpose, data, viz, periods in reports:
        p = doc.add_paragraph()
        p.add_run(name + "\n").bold = True
        p.add_run(f"Назначение: {purpose}\n")
        p.add_run(f"Данные: {data}\n")
        p.add_run(f"Визуализация: {viz}\n")
        p.add_run(f"Периоды: {periods}")
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ==================== SECTION 9: FAQ ====================
    doc.add_heading("❓ 9. Часто задаваемые вопросы", level=1)
    
    doc.add_heading("Продажи и касса", level=2)
    
    faqs_sales = [
        ("Может ли кассир продать больше товаров, чем есть на складе?", 
         "Нет. Система проверяет остатки при оформлении. Если товара недостаточно, продажа будет заблокирована."),
        ("Что произойдёт, если я отменю продажу?",
         "Продажа будет помечена как «Отменена», и весь товар автоматически вернётся на склад. Это действие нельзя отменить."),
        ("Можем ли мы сделать раздельную оплату (часть наличными, часть картой)?",
         "В данный момент нет. Каждая продажа требует один способ оплаты. Это возможное улучшение в будущем."),
        ("Как обработать возврат от покупателя?",
         'Используйте функцию "Отменить продажу" или "Возврат". Товар автоматически вернётся на склад.')
    ]
    
    for q, a in faqs_sales:
        p = doc.add_paragraph()
        p.add_run("В: " + q + "\n").bold = True
        p.add_run("О: " + a)
        doc.add_paragraph()
    
    doc.add_heading("Склад", level=2)
    
    faqs_inventory = [
        ("Что делать, если фактический остаток не совпадает с системой?",
         'Используйте функцию "Ручная корректировка остатка" (только для Админа/Менеджера). Введите правильное количество с причиной.'),
        ("Можно ли увидеть, кто изменил остаток?",
         "Да. Журнал склада показывает каждое изменение, включая пользователя, дату/время и причину."),
        ("Как узнать, когда нужно заказать товар?",
         'Установите "Минимальный остаток" для каждого товара. Когда остаток упадёт ниже этого уровня, товар появится в виджете "Низкий остаток".')
    ]
    
    for q, a in faqs_inventory:
        p = doc.add_paragraph()
        p.add_run("В: " + q + "\n").bold = True
        p.add_run("О: " + a)
        doc.add_paragraph()
    
    doc.add_heading("Поставщики и закупки", level=2)
    
    faqs_suppliers = [
        ("Можно ли принять частичную поставку от поставщика?",
         'Да. Функция "Принять товары" поддерживает частичную приёмку. Статус заказа покажет "Частично получен".'),
        ("Что если поставщик привёз больше, чем мы заказали?",
         "Система позволяет принять только до заказанного количества по каждой позиции."),
        ("Можно ли удалить заказ поставщику?",
         'Только если он в статусе "Черновик". После отправки его можно только отменить — для целей аудита.')
    ]
    
    for q, a in faqs_suppliers:
        p = doc.add_paragraph()
        p.add_run("В: " + q + "\n").bold = True
        p.add_run("О: " + a)
        doc.add_paragraph()
    
    doc.add_heading("Отчёты и аналитика", level=2)
    
    faqs_reports = [
        ("Как рассчитывается прибыль?",
         "Прибыль = Выручка (сумма продаж) - Себестоимость (закупочная цена × количество). Система автоматически рассчитывает это."),
        ("Можно ли экспортировать отчёты в Excel или PDF?",
         "В текущей версии нет. Это запланировано на будущие релизы."),
        ("Как часто обновляются данные на главной панели?",
         "В реальном времени. Каждая продажа, отмена или изменение склада мгновенно обновляет показатели.")
    ]
    
    for q, a in faqs_reports:
        p = doc.add_paragraph()
        p.add_run("В: " + q + "\n").bold = True
        p.add_run("О: " + a)
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ==================== SECTION 10: GLOSSARY ====================
    doc.add_heading("📝 10. Глоссарий", level=1)
    
    table = doc.add_table(rows=13, cols=2)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = "Термин"
    headers[1].text = "Определение"
    
    glossary = [
        ("POS", "Point of Sale (Точка продажи) — интерфейс кассы"),
        ("SKU", "Stock Keeping Unit — уникальный идентификатор товара"),
        ("Штрих-код", "Сканируемый код товара (EAN, UPC и др.)"),
        ("Себестоимость", "Цена закупки товара у поставщика"),
        ("Цена продажи", "Цена, которую платит покупатель"),
        ("Маржа", "Разница между ценой продажи и себестоимостью"),
        ("Заказ поставщику", "Заявка на закупку товаров"),
        ("COGS", "Себестоимость проданных товаров"),
        ("Низкий остаток", "Количество товара ниже настроенного минимума"),
        ("Журнал склада", "Полная история всех движений товара"),
        ("Чек", "Документ, подтверждающий продажу"),
        ("Инвентаризация", "Физический пересчёт товаров на складе")
    ]
    
    for i, (term, definition) in enumerate(glossary, 1):
        row = table.rows[i].cells
        row[0].text = term
        row[1].text = definition
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer
    footer = doc.add_paragraph()
    footer.add_run("Документ подготовлен: ").bold = True
    footer.add_run("Команда разработки\n")
    footer.add_run("По вопросам: ").bold = True
    footer.add_run("Обратитесь к системному администратору или команде разработки")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    copyright = doc.add_paragraph()
    copyright.add_run("© 2026 Sellary POS. Документ предназначен для внутреннего использования и ознакомления клиентов.")
    copyright.alignment = WD_ALIGN_PARAGRAPH.CENTER
    copyright.runs[0].font.size = Pt(9)
    copyright.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # Save document
    doc.save('BUSINESS_LOGIC_GUIDE_RU.docx')
    print("✅ Документ успешно создан: BUSINESS_LOGIC_GUIDE_RU.docx")

if __name__ == "__main__":
    create_document()
