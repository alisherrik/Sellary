# -*- coding: utf-8 -*-
"""A short, printable instruction for cashiers: the drawer count is cash only.

Written for the person closing the till, not for the manager — so it explains
the rule, shows the arithmetic on a plain example, and stops. The shop's own
figures stay in the manager's report.
"""
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ACCENT = RGBColor(0x1D, 0x4E, 0xD8)
INK = RGBColor(0x11, 0x18, 0x27)
MUTED = RGBColor(0x6B, 0x72, 0x80)
RED = RGBColor(0xB9, 0x1C, 0x1C)
GREEN = RGBColor(0x15, 0x80, 0x3D)

doc = Document()
s = doc.sections[0]
s.top_margin = s.bottom_margin = Cm(1.9)
s.left_margin = Cm(2.0)
s.right_margin = Cm(1.8)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for name, size, color, before, after in (
    ("Heading 1", 18, ACCENT, 16, 6),
    ("Heading 2", 13, INK, 12, 5),
):
    st = doc.styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = color
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)


def shade(cell, hexc):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexc)
    cell._tc.get_or_add_tcPr().append(el)


def para(text="", *, size=11, bold=False, color=None, space_after=6, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def step(text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    return p


def box(title, body, fill, tc, title_size=12, body_size=11):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.rows[0].cells[0]
    shade(cell, fill)
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.size = Pt(title_size)
    r1.font.color.rgb = tc
    p1.paragraph_format.space_after = Pt(3)
    if body:
        p2 = cell.add_paragraph()
        r2 = p2.add_run(body)
        r2.font.size = Pt(body_size)
        p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def table(headers, rows, widths=None, right=(), size=10.5, head_fill="F3F4F6"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        shade(cell, head_fill)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(0)
        if i in right:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(str(v))
            r.font.size = Pt(size)
            p.paragraph_format.space_after = Pt(0)
            if i in right:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def formula(lines):
    for line in lines:
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(11)
        r.bold = True
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.space_after = Pt(1)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ============================================================== титул =====
p = doc.add_paragraph()
r = p.add_run("В кассе — только наличные")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = ACCENT
p.paragraph_format.space_after = Pt(2)

para(
    "Почему при закрытии смены нельзя прибавлять к наличным выручку по карте",
    size=13,
    color=MUTED,
    space_after=14,
)

box(
    "Правило",
    "В поле «Посчитанные наличные в кассе» вводится только то, что физически "
    "лежит в денежном ящике: банкноты и монеты. Оплата картой, мобильный "
    "платёж и продажи в долг туда не добавляются.",
    "EFF6FF",
    ACCENT,
    title_size=13,
)

# ============================================================== 1 =========
doc.add_heading("1. Где на самом деле лежат деньги", level=1)

para(
    "Смена спрашивает у вас один простой факт: сколько наличных сейчас в ящике. "
    "Не сколько заработали, а сколько лежит. Разные виды оплаты попадают в "
    "разные места:"
)

table(
    ["Как заплатил покупатель", "Куда попали деньги", "Есть в ящике?"],
    [
        ("Наличными", "В денежный ящик кассы", "ДА"),
        ("Картой (DC, Эсхата, Alif)", "На счёт магазина в банке", "НЕТ"),
        ("Мобильным платежом", "На счёт магазина", "НЕТ"),
        ("Взял в долг", "Остались у покупателя", "НЕТ"),
    ],
    widths=[6.4, 6.4, 3.2],
)

para(
    "Поэтому в ящике никогда не лежит вся выручка смены. Там лежит только та "
    "её часть, которую отдали наличными, плюс то, что принесли в погашение "
    "долгов, минус выданные возвраты.",
)

# ============================================================== 2 =========
doc.add_heading("2. Что происходит, если прибавить карту", level=1)

para("Возьмём смену, где продали на 1 300: 1 000 наличными и 300 по карте.")

doc.add_heading("Как правильно", level=2)
formula(
    [
        "В ящике лежит:            1 000",
        "Программа ожидает:        1 000",
        "Расхождение:                  0      ← касса сошлась",
    ]
)

doc.add_heading("Как неправильно", level=2)
formula(
    [
        "В ящике лежит:            1 000",
        "Кассир вписал 1000 + 300: 1 300",
        "Программа ожидает:        1 000",
        "Расхождение:               +300      ← излишек, которого нет",
    ]
)

para(
    "Триста сомони никуда не появились — они на счёте в банке. Программа "
    "честно сообщает, что в ящике на 300 больше, чем должно быть, потому что "
    "ей сказали именно это.",
)

box(
    "И эта ошибка переходит в следующую смену",
    "Сумму, с которой закрылась смена, обычно вписывают в «На начало» "
    "следующей. Завышенная цифра переезжает дальше, и записанный остаток "
    "кассы всё сильнее расходится с реальными деньгами.",
    "FFF7ED",
    RGBColor(0xB4, 0x53, 0x09),
)

# ============================================================== 3 =========
doc.add_page_break()
doc.add_heading("3. Главная опасность: недостача исчезает", level=1)

para(
    "Излишек сам по себе выглядит безобидно. Опасно другое: он маскирует "
    "настоящую нехватку денег.",
)

para("Тот же день, но из кассы пропало 300 наличными.", bold=True, space_after=4)

formula(
    [
        "Должно быть в ящике:      1 000",
        "На самом деле в ящике:      700      ← пропало 300",
        "Кассир вписал 700 + 300:  1 000      ← прибавил карту",
        "Расхождение:                  0      ← НИЧЕГО НЕ ВИДНО",
    ]
)

box(
    "Две ошибки погасили друг друга",
    "Пропажа 300 наличными и лишние 300 по карте дали ровно ноль. Отчёт "
    "показывает, что всё в порядке. Именно ради этой проверки смена и "
    "существует — прибавляя карту, вы её выключаете.",
    "FEF2F2",
    RED,
    title_size=12,
)

para(
    "Проверка кассы защищает в первую очередь самого кассира: пока расхождение "
    "считается честно, видно, что деньги на месте. Как только в счёт "
    "подмешивают безналичные, доказать это уже нечем.",
)

# ============================================================== 4 =========
doc.add_heading("4. То же самое с продажами в долг", level=1)

para(
    "Продажа «в долг» — это товар, отданный без денег. Их в ящике нет и быть "
    "не может, поэтому в счёт они тоже не идут."
)
para(
    "А вот когда покупатель приходит и гасит долг наличными — эти деньги "
    "попадают в ящик, и программа сама их учитывает в строке «Оплата долга "
    "наличными». Отдельно прибавлять их не нужно.",
)

# ============================================================== 5 =========
doc.add_heading("5. Как закрывать смену правильно", level=1)

step("Откройте «Касса» → «Смена» и нажмите «Закрыть смену».")
step(
    "Достаньте деньги из ящика и пересчитайте их. Только банкноты и монеты, "
    "ничего больше."
)
step("Введите получившуюся сумму. Ничего к ней не прибавляйте.")
step(
    "Программа покажет: сколько ожидалось, сколько вы насчитали и разницу. "
    "Посмотрите на разницу."
)
step("Нажмите «Подтвердить закрытие»."),

box(
    "Программа предупредит",
    "Если разница окажется ровно равна выручке по карте, мобильным платежам "
    "или продажам в долг, появится жёлтое предупреждение: «Похоже, в счёт "
    "попали безналичные деньги». Увидели его — вернитесь и введите только "
    "наличные.",
    "FFFBEB",
    RGBColor(0xB4, 0x53, 0x09),
)

para(
    "Ожидаемая сумма специально не показывается, пока вы не введёте свою. "
    "Иначе её легко переписать вместо того, чтобы посчитать — и проверка "
    "снова перестанет работать.",
    color=MUTED,
    size=10.5,
)

# ============================================================== 6 =========
doc.add_heading("6. Если деньги не сошлись", level=1)

para("Расхождение — это не обвинение, а повод проверить. По порядку:")
bullet("Пересчитайте деньги ещё раз, спокойно и до конца.")
bullet(
    "Проверьте сумму «На начало»: если при открытии смены её вписали неверно, "
    "расхождение будет ровно на эту ошибку."
)
bullet(
    "Вспомните, не брали ли наличные из кассы на расходы — программа такой "
    "расход не видит, он выглядит как недостача."
)
bullet(
    "Проверьте способы оплаты в чеках: продажа по карте, пробитая как "
    "наличные, даёт недостачу; обратная ошибка — излишек."
)
bullet("Проверьте, все ли возвраты и погашения долгов проведены в программе.")

box(
    "Не подгоняйте цифру",
    "Если после проверки расхождение осталось — закройте смену с той суммой, "
    "которую действительно насчитали. Честная цифра, даже неудобная, полезнее "
    "подогнанной: по ней можно найти причину. Подогнанная не значит ничего.",
    "FEF2F2",
    RED,
)

# ============================================================== памятка ===
doc.add_page_break()
doc.add_heading("Памятка у кассы", level=1)
para("Можно распечатать и повесить рядом с денежным ящиком.", color=MUTED, size=10)

box(
    "В «Посчитано» вводится ТОЛЬКО то, что лежит в ящике",
    "",
    "EFF6FF",
    ACCENT,
    title_size=15,
)

table(
    ["Считаем", "Не считаем"],
    [
        ("Банкноты из ящика", "Оплату картой"),
        ("Монеты из ящика", "Мобильные платежи"),
        ("", "Продажи в долг"),
        ("", "Сумму «Ожидается в кассе»"),
    ],
    widths=[8.0, 8.0],
    size=12,
)

para(
    "Сначала посчитай деньги — потом смотри, что скажет программа.",
    size=13,
    bold=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    space_after=4,
)
para(
    "Если разница ровно равна выручке по карте — значит, карту посчитали "
    "как наличные.",
    size=11,
    color=RED,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    space_after=16,
)

para(
    "Sellary · инструкция по закрытию смены",
    size=9,
    color=MUTED,
    align=WD_ALIGN_PARAGRAPH.CENTER,
)

doc.save(os.environ["CARD_DOC_OUT"])
print("saved")
