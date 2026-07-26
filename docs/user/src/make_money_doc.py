# -*- coding: utf-8 -*-
"""Two new screens explained for the shop: «Деньги» and «Отчёт по закупкам».

Written for the owner and the manager, not the developer: what each screen
answers, which number means what, and the handful of mistakes that make the
figures stop agreeing.
"""
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ACCENT = RGBColor(0x1D, 0x4E, 0xD8)
INK = RGBColor(0x11, 0x18, 0x27)
MUTED = RGBColor(0x6B, 0x72, 0x80)
WARN = RGBColor(0xB4, 0x53, 0x09)
RED = RGBColor(0xB9, 0x1C, 0x1C)

doc = Document()
s = doc.sections[0]
s.top_margin = s.bottom_margin = Cm(1.9)
s.left_margin = Cm(2.1)
s.right_margin = Cm(1.8)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for name, size, color, before, after in (
    ("Heading 1", 19, ACCENT, 16, 7),
    ("Heading 2", 13.5, INK, 13, 5),
    ("Heading 3", 11.5, INK, 10, 4),
):
    st = doc.styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = color
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)


def shade(cell, hexc):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexc)
    cell._tc.get_or_add_tcPr().append(el)


def para(text="", *, size=11, bold=False, color=None, space_after=6, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def step(text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    return p


def box(title, body, fill="FFF7ED", tc=WARN, title_size=11.5):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.rows[0].cells[0]
    shade(cell, fill)
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.size = Pt(title_size)
    r1.font.color.rgb = tc
    p1.paragraph_format.space_after = Pt(3)
    if body:
        p2 = cell.add_paragraph()
        r2 = p2.add_run(body)
        r2.font.size = Pt(10.5)
        p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def table(headers, rows, widths=None, right=(), size=10):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        shade(cell, "F3F4F6")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(0)
        if i in right:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(str(v))
            r.font.size = Pt(size)
            if i == 0 and len(headers) > 1:
                r.bold = True
            p.paragraph_format.space_after = Pt(0)
            if i in right:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def formula(lines):
    for line in lines:
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(10.5)
        r.bold = True
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.space_after = Pt(1)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ================================================================ титул ====
p = doc.add_paragraph()
r = p.add_run("Деньги и закупки")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = ACCENT
p.paragraph_format.space_after = Pt(2)

para(
    "Два новых раздела: где лежат деньги магазина и что вы покупаете почём",
    size=12.5,
    color=MUTED,
    space_after=14,
)

para(
    "Раньше программа видела деньги только через продажу, возврат и погашение "
    "долга. Всё остальное — сдачу выручки в банк, оплату поставщику наличными, "
    "снятие денег с карты — записать было негде, и каждое такое движение "
    "превращалось в недостачу или излишек по кассе. Теперь для этого есть "
    "раздел «Деньги». А раздел «Отчёт по закупкам» отвечает на вопрос, который "
    "до сих пор приходилось считать в голове: что мы купили и по какой цене.",
    space_after=14,
)

# =============================================================== ЧАСТЬ 1 ===
doc.add_heading("Часть 1. Раздел «Деньги»", level=1)

doc.add_heading("Что такое счёт", level=2)
para(
    "Счёт — это место, где лежат деньги. У магазина их несколько, и программа "
    "заводит их сама:"
)
table(
    ["Счёт", "Что это", "Откуда там появляются деньги"],
    [
        ("Касса", "Денежный ящик кассира", "Продажи за наличные, погашение долгов наличными"),
        ("Банк · DC", "Счёт, куда приходят оплаты картой DC", "Каждая продажа по карте DC"),
        ("Банк · Эсхата", "То же для карты Эсхата", "Каждая продажа по карте Эсхата"),
        ("Сейф и другие", "Всё, что добавите сами", "Только то, что запишете вручную"),
    ],
    widths=[3.6, 5.4, 7.6],
)
para(
    "Продажи, возвраты и погашения долгов попадают на счета сами. Записывать "
    "вручную нужно только то, чего программа знать не может.",
)

doc.add_heading("Четыре действия", level=2)
table(
    ["Действие", "Когда", "Пример"],
    [
        (
            "Внести",
            "Деньги пришли на счёт не от продажи",
            "Владелец положил 500 в кассу на размен",
        ),
        (
            "Изъять",
            "Деньги ушли со счёта",
            "Отдали 150 поставщику за хлеб · сдали 5 000 в банк",
        ),
        (
            "Перевод",
            "Деньги перешли с одного счёта на другой",
            "Сняли 300 с карты DC и положили в кассу",
        ),
        (
            "Сверить",
            "Пересчитали и в реальности другая сумма",
            "В сейфе 940, а программа показывает 1 000",
        ),
    ],
    widths=[2.8, 5.6, 8.2],
)

box(
    "Перевод — не «внести» и не «изъять»",
    "Когда деньги снимают с карты и кладут в кассу, они не появляются и не "
    "пропадают — они переезжают. Поэтому это одно действие «Перевод», а не два "
    "отдельных. Программа запишет обе стороны сразу: на банке станет меньше, в "
    "кассе больше, а общая сумма денег магазина не изменится.",
    fill="EFF6FF",
    tc=ACCENT,
)

doc.add_heading("Как это выглядит", level=3)
formula(
    [
        "ДО:    Касса 12 583.12    Банк · DC 4 120.00    Всего 16 703.12",
        "Перевод 300 из банка в кассу",
        "ПОСЛЕ: Касса 12 883.12    Банк · DC 3 820.00    Всего 16 703.12",
    ]
)
para(
    "Общая сумма та же — деньги просто в другом месте. Именно так и должно "
    "быть.",
    space_after=10,
)

doc.add_heading("«Сверить» вместо исправления", level=2)
para(
    "Записи о движении денег нельзя отредактировать и нельзя удалить. Если "
    "ошиблись — пересчитайте счёт и нажмите «Сверить», указав, сколько на нём "
    "на самом деле. Программа сама запишет разницу отдельной строкой."
)
para(
    "Так в истории остаётся правда: видно и первоначальную запись, и то, что "
    "её поправили, и когда. Исправленная задним числом касса не стоит ничего — "
    "по ней уже нельзя понять, что произошло.",
    space_after=8,
)

doc.add_heading("Связь со сменой", level=2)
para(
    "Движения по кассе можно записать прямо со страницы смены — там появились "
    "кнопки «Внести» и «Изъять». Они сразу входят в расчёт:"
)
formula(
    [
        "На начало  +  Продажи наличными  +  Оплата долга наличными",
        "           +  Внесения  −  Возвраты наличными  −  Изъятия",
        "           =  Ожидается в кассе",
    ]
)
para(
    "Каждое движение видно в блоке «Прочие движения денег» с причиной и "
    "комментарием. В выручку они не входят: взяли деньги из кассы на хлеб — "
    "продали от этого не меньше и не больше.",
)

box(
    "Движение по кассе записывается только при открытой смене",
    "Итоги закрытой смены заморожены и не пересчитываются, поэтому задним "
    "числом добавить в неё движение нельзя. Записывайте расход в тот же день, "
    "когда взяли деньги.",
)

doc.add_heading("Что «Деньги» НЕ меняют", level=2)
para(
    "Ни одна запись в этом разделе не влияет на выручку, прибыль и отчёты по "
    "продажам. Сдача выручки в банк не уменьшает выручку — деньги просто "
    "переехали. Оплата поставщику наличными не уменьшает прибыль сегодняшнего "
    "дня — стоимость товара учитывается тогда, когда товар продан.",
    space_after=8,
)

# =============================================================== ЧАСТЬ 2 ===
doc.add_page_break()
doc.add_heading("Часть 2. Отчёт по закупкам", level=1)

para(
    "Раздел «Закупки» → «Отчёт по закупкам». Отвечает на четыре вопроса: "
    "сколько потратили на товар, что именно купили и почём, у кого закупались, "
    "и что заказали, но ещё не получили."
)

box(
    "Считается принятый товар, а не заказанный",
    "Пока заказ не принят на склад, он ничего не стоит: товара нет, деньги не "
    "потрачены. Такие заказы вынесены на отдельную вкладку «Заказано, но не "
    "пришло» и ни в одну сумму отчёта не входят.",
    fill="EFF6FF",
    tc=ACCENT,
)

doc.add_heading("Четыре цифры наверху", level=2)
table(
    ["Показатель", "Что означает"],
    [
        ("Потрачено на товар", "Сумма всего принятого товара за период."),
        ("Средняя поставка", "Потрачено ÷ число поставок. Насколько крупными партиями берёте."),
        ("Наименований", "Сколько разных товаров закупали."),
        ("Поставщиков", "У скольких закупались."),
    ],
    widths=[4.6, 12.0],
)

doc.add_heading("Вкладка «По товарам»", level=2)
para("Главная таблица: что купили и по какой цене.")
table(
    ["Столбец", "Что означает"],
    [
        ("Куплено", "Сколько единиц приняли за период."),
        ("Потрачено", "Сколько денег ушло на этот товар."),
        ("Доля", "Какую часть всех закупок он занимает. Показывает, на что уходят деньги."),
        (
            "Средняя цена",
            "Потрачено ÷ куплено. Средняя цена закупки, если брали по разной цене.",
        ),
        (
            "Было → стало",
            "Цена в первой и в последней поставке за период. Направление, куда движется цена.",
        ),
        (
            "Изменение",
            "На сколько процентов выросла или упала закупочная цена. Рост красным — "
            "он съедает прибыль.",
        ),
        ("Продаём по", "Текущая цена продажи — сразу видно, осталась ли наценка."),
        ("Поставок", "Сколько раз привозили этот товар."),
    ],
    widths=[3.6, 13.0],
)

box(
    "Зачем «было → стало», если есть средняя цена",
    "Средняя цена скрывает направление. Если сахар привезли по 10, а потом по "
    "12.50, средняя будет 11.67 — и по ней не видно, что цена растёт. Столбец "
    "«Было → стало» показывает 10.00 → 12.50, а «Изменение» — +25 %. Это и "
    "есть сигнал: пора либо поднимать цену продажи, либо искать другого "
    "поставщика.",
    fill="EFF6FF",
    tc=ACCENT,
)

doc.add_heading("Пример строки", level=3)
table(
    ["Товар", "Куплено", "Потрачено", "Средняя", "Было → стало", "Изменение", "Продаём по"],
    [("Сахар 1кг", "30 кг", "350.00", "11.6667", "10.00 → 12.50", "+25.0 %", "18.00")],
    widths=[3.0, 2.0, 2.2, 2.2, 3.0, 2.0, 2.2],
    right=(1, 2, 3, 4, 5, 6),
    size=9.5,
)
para(
    "Читается так: за период приняли 30 кг сахара на 350 сомони. В среднем "
    "по 11.67, но цена выросла с 10.00 до 12.50 — на четверть. Продаём по "
    "18.00: наценка была 8.00, стала 5.50.",
    space_after=10,
)

doc.add_heading("Вкладка «По поставщикам»", level=2)
para(
    "Кому и сколько заплатили, какую долю закупок занимает каждый и когда была "
    "последняя поставка. Полезно перед разговором о цене: поставщик, на "
    "которого приходится половина закупок, обычно готов обсуждать скидку."
)

doc.add_heading("Вкладка «Заказано, но не пришло»", level=2)
para(
    "Заказы, которые отправили поставщику, но приняли не полностью. Показывает "
    "сумму заказа и сколько позиций ещё не привезли. Эти деньги пока не "
    "потрачены — товара нет."
)

doc.add_heading("Период", level=2)
para(
    "Кнопки «7 дней», «30 дней», «90 дней», «Год» вверху справа. День "
    "закрывается по местному времени магазина, поэтому вечерняя поставка "
    "попадает в тот день, когда её действительно приняли."
)

# ================================================================ финал ====
doc.add_heading("Короткая памятка", level=1)
bullet("Взяли деньги из кассы — сразу нажмите «Изъять» и выберите причину.")
bullet("Положили деньги в кассу не от продажи — «Внести».")
bullet("Сняли с карты и положили в кассу — «Перевод», а не «Внести».")
bullet("Ошиблись — не исправляйте запись, нажмите «Сверить».")
bullet("Движения по кассе записываются только при открытой смене.")
bullet("Ни одно движение денег не меняет выручку и прибыль.")
bullet("В отчёте по закупкам считается принятый товар, а не заказанный.")
bullet("Смотрите столбец «Изменение»: там видно, что дорожает.")

doc.add_paragraph()
para(
    "Sellary · разделы «Деньги» и «Отчёт по закупкам»",
    size=9,
    color=MUTED,
    align=WD_ALIGN_PARAGRAPH.CENTER,
)

doc.save(os.environ["MONEY_DOC_OUT"])
print("saved")
